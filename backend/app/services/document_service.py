import os
import fitz  # PyMuPDF
import docx
import openpyxl

class DocumentService:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Detects file type and extracts raw text.
        Supports: PDF, DOCX, XLSX, TXT
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path.lower())
        
        try:
            if ext == ".pdf":
                return DocumentService._extract_pdf(file_path)
            elif ext == ".docx":
                return DocumentService._extract_docx(file_path)
            elif ext == ".xlsx" or ext == ".xls":
                return DocumentService._extract_xlsx(file_path)
            elif ext in [".txt", ".csv", ".json"]:
                return DocumentService._extract_txt(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {ext}")
        except Exception as e:
            # Re-raise with a clear message
            raise RuntimeError(f"Failed to extract text from {ext} file: {str(e)}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        text_content = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_content.append(page.get_text())
        
        extracted = "\n".join(text_content).strip()
        if not extracted:
            raise ValueError("No extractable text found. PDF might be scanned/image-only.")
        return extracted

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
        extracted = "\n".join(text_content).strip()
        if not extracted:
            raise ValueError("No extractable text found in DOCX file.")
        return extracted

    @staticmethod
    def _extract_xlsx(file_path: str) -> str:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_content = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_content.append(f"--- Sheet: {sheet_name} ---")
            for row in sheet.iter_rows(values_only=True):
                # Format non-empty row items
                row_str = " | ".join(str(cell).strip() for cell in row if cell is not None)
                if row_str.strip():
                    text_content.append(row_str)
                    
        extracted = "\n".join(text_content).strip()
        if not extracted:
            raise ValueError("No extractable text found in XLSX spreadsheet.")
        return extracted

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
